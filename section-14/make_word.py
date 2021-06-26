import docx

document = docx.Document()
document.add_paragraph("Hello, this is a paragraph.")
document.add_paragraph("This is another paragraph!")

paragraph_1 = document.paragraphs[0]
to_bold = paragraph_1.add_run("This is a new run.")
to_bold.bold = True
# paragraph_1.runs[1]

document.save("files/section-14/make_word.docx")